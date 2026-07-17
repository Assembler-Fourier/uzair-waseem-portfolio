from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from pypdf import PdfReader

from cv_engine import OUTPUT_ROOT, PUBLIC_DIR, PUBLIC_VARIANT_DIR, load_profile


PRIVATE_PATTERNS = [r"899739932", r"89\s*973\s*9932", r"K78\s*H9Y9", r"15\s+Aderrig"]
REQUIRED_SECTIONS = ["PROFILE", "TECHNICAL SKILLS", "PROFESSIONAL EXPERIENCE", "SELECTED ENGINEERING PROJECTS", "EDUCATION"]


def pdf_text(path: Path) -> tuple[PdfReader, str]:
    reader = PdfReader(str(path))
    return reader, "\n".join(page.extract_text() or "" for page in reader.pages)


def main() -> None:
    profile = load_profile()
    failures: list[str] = []
    fingerprints: dict[str, str] = {}
    results = []

    for track, track_data in profile["tracks"].items():
        public_path = PUBLIC_VARIANT_DIR / track_data["publicFilename"]
        reader, text = pdf_text(public_path)
        if len(reader.pages) != 1:
            failures.append(f"{track}: expected one page, found {len(reader.pages)}")
        if not 2500 <= len(text) <= 5200:
            failures.append(f"{track}: unexpected extracted text length {len(text)}")
        link_annotations = sum(
            annotation.get_object().get("/Subtype") == "/Link"
            for page in reader.pages
            for annotation in (page.get("/Annots") or [])
        )
        if link_annotations < 8:
            failures.append(f"{track}: expected working PDF links, found {link_annotations}")
        for section in REQUIRED_SECTIONS:
            if section not in text:
                failures.append(f"{track}: missing section {section}")
        if any(re.search(pattern, text, re.IGNORECASE) for pattern in PRIVATE_PATTERNS):
            failures.append(f"{track}: public CV contains a private phone/address pattern")
        if "Pre-Launch" in text or "pre-launch" in text:
            failures.append(f"{track}: stale Theory Coach pre-launch label")
        if any(project_id == "theory" for project_id in track_data["projectOrder"]):
            if "Irish Theory Test Coach - Live Deployed EdTech Platform" not in text:
                failures.append(f"{track}: live Theory Coach title missing")
            if "irishtheorycoach.ie" not in text:
                failures.append(f"{track}: canonical Theory Coach URL missing")
        fingerprints[track] = hashlib.sha256(text.encode("utf-8")).hexdigest()
        output_stem = Path(track_data["outputFilename"]).stem
        docx_path = OUTPUT_ROOT / "cv" / track / f"{output_stem}.docx"
        if not docx_path.exists():
            failures.append(f"{track}: generated DOCX is missing")
        else:
            document = Document(docx_path)
            docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            if document.tables:
                failures.append(f"{track}: DOCX contains tables")
            for section in REQUIRED_SECTIONS:
                if section not in docx_text:
                    failures.append(f"{track}: DOCX missing section {section}")
            if any(re.search(pattern, docx_text, re.IGNORECASE) for pattern in PRIVATE_PATTERNS):
                failures.append(f"{track}: DOCX contains a private phone/address pattern")
            bullet_paragraphs = [
                paragraph for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name == "CV Bullet"
            ]
            if len(bullet_paragraphs) < 10:
                failures.append(f"{track}: DOCX has too few evidence bullets")
            if any(paragraph._p.pPr is None or paragraph._p.pPr.numPr is None for paragraph in bullet_paragraphs):
                failures.append(f"{track}: DOCX bullet is not backed by Word numbering")
            hyperlink_count = sum(
                relation.reltype == RELATIONSHIP_TYPE.HYPERLINK
                for relation in document.part.rels.values()
            )
            if hyperlink_count < 8:
                failures.append(f"{track}: DOCX has too few working hyperlink relationships")
        results.append(
            {
                "track": track,
                "pages": len(reader.pages),
                "textCharacters": len(text),
                "pdfLinks": link_annotations,
                "pdf": str(public_path),
                "docx": str(docx_path),
            }
        )

    if len(set(fingerprints.values())) != len(fingerprints):
        failures.append("Role-specific public CVs are not all distinct")

    master_reader, master_text = pdf_text(PUBLIC_DIR / "Uzair-Waseem-CV.pdf")
    software_reader, software_text = pdf_text(PUBLIC_VARIANT_DIR / profile["tracks"]["software"]["publicFilename"])
    if master_text != software_text or len(master_reader.pages) != len(software_reader.pages):
        failures.append("Master public CV does not match the software track")

    report_path = OUTPUT_ROOT / "cv-validation.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps({"results": results, "failures": failures}, indent=2) + "\n", encoding="utf-8")
    if failures:
        raise SystemExit("\n".join(failures))
    print(f"Validated {len(results)} distinct one-page public CVs")
    print(report_path)


if __name__ == "__main__":
    main()
