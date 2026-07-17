from __future__ import annotations

import copy
import json
import os
import re
import shutil
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
PROFILE_PATH = ROOT / "cv" / "profile.json"
OUTPUT_ROOT = ROOT / "output"
PUBLIC_DIR = ROOT / "public"
PUBLIC_VARIANT_DIR = PUBLIC_DIR / "cv"
MASTER_FILENAME = "Uzair-Waseem-CV.pdf"
TRACK_PRIORITY = ["software", "full-stack", "backend", "qa", "security"]


def load_profile() -> dict[str, Any]:
    return json.loads(PROFILE_PATH.read_text(encoding="utf-8"))


def slugify(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9]+", "-", value.strip()).strip("-")
    return value or "Target-Role"


def term_pattern(term: str) -> re.Pattern[str]:
    return re.compile(rf"(?<![A-Za-z0-9]){re.escape(term.lower())}(?![A-Za-z0-9])")


def term_count(text: str, terms: list[str]) -> int:
    lowered = text.lower()
    return sum(len(term_pattern(term).findall(lowered)) for term in terms)


def score_tracks(profile: dict[str, Any], job_text: str) -> dict[str, int]:
    lowered = job_text.lower()
    scores: dict[str, int] = {}
    for key, track in profile["tracks"].items():
        scores[key] = sum(term_count(lowered, [signal]) * 3 for signal in track["signals"])

    explicit_boosts = {
        "qa": ["qa automation", "quality assurance", "sdet", "test automation engineer"],
        "backend": ["backend engineer", "back-end engineer", "api engineer"],
        "full-stack": ["full-stack engineer", "full stack engineer", "fullstack engineer"],
        "security": ["application security", "secure software", "security-aware software"],
        "software": ["software engineer", "graduate software engineer", "junior software engineer"],
    }
    for key, phrases in explicit_boosts.items():
        if any(term_count(lowered, [phrase]) for phrase in phrases):
            scores[key] += 20
    return scores


def select_track(profile: dict[str, Any], job_text: str, requested_track: str | None = None) -> tuple[str, dict[str, int]]:
    if requested_track:
        if requested_track not in profile["tracks"]:
            choices = ", ".join(sorted(profile["tracks"]))
            raise ValueError(f"Unknown track '{requested_track}'. Choose from: {choices}")
        return requested_track, score_tracks(profile, job_text)

    scores = score_tracks(profile, job_text)
    priority = {key: len(TRACK_PRIORITY) - index for index, key in enumerate(TRACK_PRIORITY)}
    selected = max(scores, key=lambda key: (scores[key], priority.get(key, 0)))
    return selected, scores


def match_verified_skills(profile: dict[str, Any], job_text: str) -> list[str]:
    matches = []
    for skill, aliases in profile["skills"].items():
        count = term_count(job_text, aliases)
        if count:
            matches.append((skill, count))
    matches.sort(key=lambda item: (-item[1], list(profile["skills"]).index(item[0])))
    return [skill for skill, _ in matches]


def match_unsupported_terms(profile: dict[str, Any], job_text: str) -> list[str]:
    return [
        label
        for label, aliases in profile["unsupportedTerms"].items()
        if term_count(job_text, aliases)
    ]


def reorder_skill_groups(track: dict[str, Any], matched_skills: list[str]) -> list[list[Any]]:
    matched = set(matched_skills)
    groups = []
    for label, items in track["skillGroups"]:
        reordered = [item for item in items if item in matched]
        reordered.extend(item for item in items if item not in matched)
        groups.append([label, reordered])
    return groups


def reorder_experience(profile: dict[str, Any], matched_skills: list[str]) -> list[dict[str, Any]]:
    matched = {item.lower() for item in matched_skills}
    experience = copy.deepcopy(profile["experience"])
    for role in experience:
        indexed = list(enumerate(role["bullets"]))
        indexed.sort(
            key=lambda item: (
                -sum(tag.lower() in matched for tag in item[1]["tags"]),
                item[0],
            )
        )
        role["bullets"] = [item for _, item in indexed]
    return experience


def rank_projects(
    profile: dict[str, Any],
    track: dict[str, Any],
    job_text: str,
    matched_skills: list[str],
) -> list[str]:
    base_order = track["projectOrder"]
    matched = {item.lower() for item in matched_skills}
    scores = {}
    for project_id, project in profile["projects"].items():
        base_index = base_order.index(project_id) if project_id in base_order else len(base_order) + 2
        score = 40 - (base_index * 4)
        score += sum(5 for tag in project["tags"] if tag.lower() in matched)
        score += sum(2 for tag in project["tags"] if term_count(job_text, [tag]))
        scores[project_id] = score
    stable_order = list(profile["projects"])
    return sorted(scores, key=lambda key: (-scores[key], stable_order.index(key)))[:3]


def build_variant(
    profile: dict[str, Any],
    job_text: str = "",
    requested_track: str | None = None,
    target_title: str = "",
    company: str = "",
) -> tuple[dict[str, Any], dict[str, Any]]:
    selected_track, track_scores = select_track(profile, job_text, requested_track)
    track = copy.deepcopy(profile["tracks"][selected_track])
    matched_skills = match_verified_skills(profile, job_text)
    unsupported = match_unsupported_terms(profile, job_text)
    project_ids = rank_projects(profile, track, job_text, matched_skills)

    variant = {
        **track,
        "track": selected_track,
        "skillGroups": reorder_skill_groups(track, matched_skills),
        "projects": [copy.deepcopy(profile["projects"][project_id]) for project_id in project_ids],
        "experience": reorder_experience(profile, matched_skills),
        "education": copy.deepcopy(profile["education"]),
        "candidate": copy.deepcopy(profile["candidate"]),
    }
    report = {
        "schemaVersion": 1,
        "target": {"title": target_title or None, "company": company or None},
        "selectedTrack": selected_track,
        "trackScores": track_scores,
        "matchedVerifiedSkills": matched_skills,
        "unsupportedRequirementsNotAdded": unsupported,
        "selectedProjects": [profile["projects"][project_id]["title"] for project_id in project_ids],
        "factCheckRequired": profile["factCheck"],
        "policy": "The tailor may reorder verified evidence but cannot invent skills, dates, metrics or outcomes."
    }
    return variant, report


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=24.5,
            leading=26,
            textColor=colors.HexColor("#111827"),
            spaceAfter=2.0,
        ),
        "headline": ParagraphStyle(
            "Headline",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.35,
            leading=12.05,
            textColor=colors.HexColor("#0f766e"),
            spaceAfter=3.2,
        ),
        "contact": ParagraphStyle(
            "Contact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.75,
            leading=10.35,
            textColor=colors.HexColor("#374151"),
            spaceAfter=1.1,
        ),
        "section": ParagraphStyle(
            "Section",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.25,
            leading=11.75,
            textColor=colors.HexColor("#0f766e"),
            spaceBefore=6.2,
            spaceAfter=2.3,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.15,
            leading=11.15,
            textColor=colors.HexColor("#111827"),
            spaceAfter=1.8,
        ),
        "item": ParagraphStyle(
            "Item",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=11.15,
            textColor=colors.HexColor("#111827"),
            spaceBefore=2.4,
            spaceAfter=0.5,
        ),
        "meta": ParagraphStyle(
            "Meta",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.25,
            leading=9.6,
            textColor=colors.HexColor("#4b5563"),
            spaceAfter=0.55,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.95,
            leading=10.75,
            leftIndent=8,
            firstLineIndent=-8,
            textColor=colors.HexColor("#111827"),
            spaceAfter=0.9,
        ),
        "link": ParagraphStyle(
            "Link",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.3,
            textColor=colors.HexColor("#2563eb"),
            spaceAfter=0.55,
        ),
    }


def pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def add_pdf_section(story: list[Any], title: str, styles: dict[str, ParagraphStyle]) -> None:
    story.append(pdf_paragraph(title.upper(), styles["section"]))
    story.append(
        HRFlowable(
            width="100%",
            thickness=0.45,
            color=colors.HexColor("#cbd5d1"),
            spaceBefore=0,
            spaceAfter=1.5,
        )
    )


def pdf_links(links: list[list[str]], styles: dict[str, ParagraphStyle]) -> Paragraph:
    content = " | ".join(
        f'<link href="{escape(url)}" color="#2563eb">{escape(label)}</link>'
        for label, url in links
    )
    return pdf_paragraph(content, styles["link"])


def contact_text(candidate: dict[str, str], phone: str | None = None) -> str:
    parts = [escape(candidate["location"])]
    if phone:
        parts.append(escape(phone))
    parts.append(f'<link href="mailto:{escape(candidate["email"])}" color="#2563eb">{escape(candidate["email"])}</link>')
    return " | ".join(parts)


def build_pdf_story(variant: dict[str, Any], phone: str | None = None) -> list[Any]:
    styles = pdf_styles()
    candidate = variant["candidate"]
    story: list[Any] = [
        pdf_paragraph(escape(candidate["name"]), styles["name"]),
        pdf_paragraph(escape(variant["headline"]), styles["headline"]),
        pdf_paragraph(contact_text(candidate, phone), styles["contact"]),
        pdf_paragraph(
            f'<link href="{candidate["portfolio"]}" color="#2563eb">uzairwaseem.com</link> | '
            f'<link href="{candidate["linkedin"]}" color="#2563eb">linkedin.com/in/uzair-waseem</link> | '
            f'<link href="{candidate["github"]}" color="#2563eb">github.com/Assembler-Fourier</link>',
            styles["contact"],
        ),
    ]

    add_pdf_section(story, "Profile", styles)
    story.append(pdf_paragraph(escape(variant["summary"]), styles["body"]))

    add_pdf_section(story, "Technical Skills", styles)
    for label, items in variant["skillGroups"]:
        story.append(pdf_paragraph(f"<b>{escape(label)}:</b> {escape(', '.join(items))}", styles["body"]))

    add_pdf_section(story, "Professional Experience", styles)
    for role in variant["experience"]:
        block = [
            pdf_paragraph(f"{escape(role['title'])} - {escape(role['company'])}", styles["item"]),
            pdf_paragraph(f"{escape(role['location'])} | {escape(role['period'])} | {escape(role['url'].replace('https://', '').rstrip('/'))}", styles["meta"]),
        ]
        block.extend(pdf_paragraph(f"- {escape(bullet['text'])}", styles["bullet"]) for bullet in role["bullets"])
        story.append(KeepTogether(block))

    add_pdf_section(story, "Selected Engineering Projects", styles)
    for project in variant["projects"]:
        block = [
            pdf_paragraph(escape(project["title"]), styles["item"]),
            pdf_paragraph(f"<b>Stack:</b> {escape(project['stack'])}", styles["meta"]),
            pdf_links(project["links"], styles),
        ]
        block.extend(pdf_paragraph(f"- {escape(bullet)}", styles["bullet"]) for bullet in project["bullets"])
        story.append(KeepTogether(block))

    add_pdf_section(story, "Education", styles)
    for education in variant["education"]:
        story.append(
            pdf_paragraph(
                f"<b>{escape(education['degree'])}</b> - {escape(education['school'])} | {escape(education['period'])}",
                styles["body"],
            )
        )
    story.append(Spacer(1, 2.5))
    return story


def draw_pdf_footer(canvas: Canvas, document: SimpleDocTemplate) -> None:
    canvas.saveState()
    width, _ = A4
    y = 9.5 * mm
    canvas.setStrokeColor(colors.HexColor("#cbd5d1"))
    canvas.setLineWidth(0.45)
    canvas.line(document.leftMargin, y + 7, width - document.rightMargin, y + 7)
    canvas.setFont("Helvetica", 7.1)
    canvas.setFillColor(colors.HexColor("#4b5563"))
    canvas.drawString(document.leftMargin, y - 1, "Dublin, Ireland | Graduate and junior engineering roles")
    canvas.drawRightString(width - document.rightMargin, y - 1, "uzairwaseem.com")
    canvas.restoreState()


def build_pdf(variant: dict[str, Any], output_path: Path, phone: str | None = None) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=11 * mm,
        bottomMargin=16 * mm,
        title=f"Uzair Waseem - {variant['headline'].split('|')[0].strip().title()} CV",
        author="Uzair Waseem",
        subject=variant["subject"],
        keywords=variant["keywords"],
    )
    document.build(build_pdf_story(variant, phone), onFirstPage=draw_pdf_footer, onLaterPages=draw_pdf_footer)
    return output_path


def set_cell_free_section_defaults(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(15.5)
    section.right_margin = Mm(15.5)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(13)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(5)


def set_run_font(run, size: float, bold: bool = False, color: str = "111827") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), "Arial")


def add_docx_hyperlink(paragraph, text: str, url: str, size: float = 8.25) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Arial")
    run_fonts.set(qn("w:hAnsi"), "Arial")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.extend([color, underline, run_fonts, font_size])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor(17, 24, 39)
    normal.paragraph_format.space_after = Pt(1.4)
    normal.paragraph_format.line_spacing = 1.03

    definitions = {
        "CV Name": (24.0, True, "111827", 0, 1),
        "CV Headline": (10.5, True, "0F766E", 0, 2.2),
        "CV Contact": (8.6, False, "374151", 0, 0.7),
        "CV Section": (10.2, True, "0F766E", 5.2, 1.8),
        "CV Item": (9.55, True, "111827", 2.0, 0.3),
        "CV Meta": (8.25, False, "4B5563", 0, 0.45),
        "CV Body": (9.25, False, "111827", 0, 1.2),
        "CV Link": (8.25, False, "2563EB", 0, 0.55),
    }
    for name, (size, bold, color, before, after) in definitions.items():
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH) if name not in styles else styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.03
        style.paragraph_format.keep_with_next = name in {"CV Name", "CV Headline", "CV Section", "CV Item"}

    bullet = styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH) if "CV Bullet" not in styles else styles["CV Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.05)
    bullet.font.color.rgb = RGBColor(17, 24, 39)
    bullet.paragraph_format.space_after = Pt(1.0)
    bullet.paragraph_format.line_spacing = 1.03


def create_docx_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "\u2022"), ("w:lvlJc", "left")):
        child = OxmlElement(tag)
        child.set(qn("w:val"), value)
        level.append(child)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indent])
    level.append(paragraph_properties)
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Arial")
    run_fonts.set(qn("w:hAnsi"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    run_properties.extend([run_fonts, color])
    level.append(run_properties)
    abstract.append(level)
    first_concrete = numbering.find(qn("w:num"))
    if first_concrete is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_concrete), abstract)

    concrete = OxmlElement("w:num")
    concrete.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    concrete.append(reference)
    numbering.append(concrete)
    return num_id


def add_docx_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph(style="CV Bullet")
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numbering_properties.extend([level, number])
    paragraph_properties.append(numbering_properties)
    value = paragraph.add_run(text)
    set_run_font(value, 9.05)


def add_docx_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="CV Section")
    paragraph.add_run(title.upper())


def add_docx_text(document: Document, text: str, style: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, 9.25, bold=True)
        suffix = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(suffix, 9.25)
    else:
        paragraph.add_run(text)


def build_docx(variant: dict[str, Any], output_path: Path, phone: str | None = None) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_cell_free_section_defaults(document)
    configure_docx_styles(document)
    bullet_num_id = create_docx_bullet_numbering(document)
    candidate = variant["candidate"]

    document.core_properties.title = f"Uzair Waseem - {variant['headline'].split('|')[0].strip().title()} CV"
    document.core_properties.subject = variant["subject"]
    document.core_properties.author = "Uzair Waseem"
    document.core_properties.keywords = variant["keywords"]

    document.add_paragraph(candidate["name"], style="CV Name")
    document.add_paragraph(variant["headline"], style="CV Headline")
    contact = document.add_paragraph(style="CV Contact")
    contact.add_run(candidate["location"])
    if phone:
        contact.add_run(f" | {phone}")
    contact.add_run(" | ")
    add_docx_hyperlink(contact, candidate["email"], f"mailto:{candidate['email']}", size=8.6)

    links = document.add_paragraph(style="CV Contact")
    add_docx_hyperlink(links, "uzairwaseem.com", candidate["portfolio"], size=8.6)
    links.add_run(" | ")
    add_docx_hyperlink(links, "linkedin.com/in/uzair-waseem", candidate["linkedin"], size=8.6)
    links.add_run(" | ")
    add_docx_hyperlink(links, "github.com/Assembler-Fourier", candidate["github"], size=8.6)

    add_docx_section(document, "Profile")
    document.add_paragraph(variant["summary"], style="CV Body")

    add_docx_section(document, "Technical Skills")
    for label, items in variant["skillGroups"]:
        paragraph = document.add_paragraph(style="CV Body")
        prefix = paragraph.add_run(f"{label}: ")
        set_run_font(prefix, 9.25, bold=True)
        value = paragraph.add_run(", ".join(items))
        set_run_font(value, 9.25)

    add_docx_section(document, "Professional Experience")
    for role in variant["experience"]:
        document.add_paragraph(f"{role['title']} - {role['company']}", style="CV Item")
        document.add_paragraph(
            f"{role['location']} | {role['period']} | {role['url'].replace('https://', '').rstrip('/')}",
            style="CV Meta",
        )
        for bullet in role["bullets"]:
            add_docx_bullet(document, bullet["text"], bullet_num_id)

    add_docx_section(document, "Selected Engineering Projects")
    for project in variant["projects"]:
        document.add_paragraph(project["title"], style="CV Item")
        stack = document.add_paragraph(style="CV Meta")
        stack_prefix = stack.add_run("Stack: ")
        set_run_font(stack_prefix, 8.25, bold=True, color="4B5563")
        stack_value = stack.add_run(project["stack"])
        set_run_font(stack_value, 8.25, color="4B5563")
        link_paragraph = document.add_paragraph(style="CV Link")
        for index, (label, url) in enumerate(project["links"]):
            if index:
                link_paragraph.add_run(" | ")
            add_docx_hyperlink(link_paragraph, label, url, size=8.25)
        for bullet in project["bullets"]:
            add_docx_bullet(document, bullet, bullet_num_id)

    add_docx_section(document, "Education")
    for education in variant["education"]:
        paragraph = document.add_paragraph(style="CV Body")
        degree = paragraph.add_run(education["degree"])
        set_run_font(degree, 9.25, bold=True)
        rest = paragraph.add_run(f" - {education['school']} | {education['period']}")
        set_run_font(rest, 9.25)

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Dublin, Ireland | Graduate and junior engineering roles | uzairwaseem.com")
    set_run_font(footer_run, 7.4, color="4B5563")

    document.save(output_path)
    return output_path


def build_plain_text(variant: dict[str, Any], phone: str | None = None) -> str:
    candidate = variant["candidate"]
    contact = [candidate["location"]]
    if phone:
        contact.append(phone)
    contact.append(candidate["email"])
    lines = [
        candidate["name"],
        variant["headline"],
        " | ".join(contact),
        f"Portfolio: {candidate['portfolio']}",
        f"LinkedIn: {candidate['linkedin']}",
        f"GitHub: {candidate['github']}",
        "",
        "PROFILE",
        variant["summary"],
        "",
        "TECHNICAL SKILLS",
    ]
    lines.extend(f"{label}: {', '.join(items)}" for label, items in variant["skillGroups"])
    lines.extend(["", "PROFESSIONAL EXPERIENCE"])
    for role in variant["experience"]:
        lines.extend(
            [
                f"{role['title']} - {role['company']}",
                f"{role['location']} | {role['period']} | {role['url']}",
            ]
        )
        lines.extend(f"- {bullet['text']}" for bullet in role["bullets"])
        lines.append("")
    lines.append("SELECTED ENGINEERING PROJECTS")
    for project in variant["projects"]:
        lines.extend([project["title"], f"Stack: {project['stack']}"])
        lines.extend(f"{label}: {url}" for label, url in project["links"])
        lines.extend(f"- {bullet}" for bullet in project["bullets"])
        lines.append("")
    lines.append("EDUCATION")
    lines.extend(
        f"{education['degree']} - {education['school']} | {education['period']}"
        for education in variant["education"]
    )
    return "\n".join(lines).strip() + "\n"


def write_report(report: dict[str, Any], output_path: Path, phone_included: bool) -> Path:
    payload = {**report, "privatePhoneIncluded": phone_included}
    output_path.write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")
    return output_path


def write_application_notes(report: dict[str, Any], output_path: Path) -> Path:
    matched = ", ".join(report["matchedVerifiedSkills"]) or "No explicit verified technology matches detected"
    unsupported = ", ".join(report["unsupportedRequirementsNotAdded"]) or "None detected"
    projects = "\n".join(f"- {project}" for project in report["selectedProjects"])
    fact_checks = "\n".join(f"- {item}" for item in report["factCheckRequired"])
    target = report["target"]
    output_path.write_text(
        "\n".join(
            [
                "# Tailored CV Review",
                "",
                f"Target role: {target['title'] or 'Not supplied'}",
                f"Target company: {target['company'] or 'Not supplied'}",
                f"Selected track: {report['selectedTrack']}",
                "",
                "## Verified Matches",
                matched,
                "",
                "## Requirements Not Added",
                unsupported,
                "",
                "## Selected Project Order",
                projects,
                "",
                "## Personal Fact Check",
                fact_checks,
                "",
                "The tailor reorders verified evidence only. Read every line before submitting.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    return output_path


def build_package(
    variant: dict[str, Any],
    report: dict[str, Any],
    output_dir: Path,
    base_name: str,
    phone: str | None = None,
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "pdf": build_pdf(variant, output_dir / f"{base_name}.pdf", phone),
        "docx": build_docx(variant, output_dir / f"{base_name}.docx", phone),
        "text": output_dir / f"{base_name}.txt",
        "report": output_dir / f"{base_name}-match-report.json",
        "notes": output_dir / f"{base_name}-review.md",
    }
    paths["text"].write_text(build_plain_text(variant, phone), encoding="utf-8")
    write_report(report, paths["report"], bool(phone))
    write_application_notes(report, paths["notes"])
    return paths


def generate_base_tracks(track_keys: list[str] | None = None, publish: bool = True) -> dict[str, dict[str, Path]]:
    profile = load_profile()
    selected = track_keys or list(profile["tracks"])
    results: dict[str, dict[str, Path]] = {}
    for key in selected:
        variant, report = build_variant(profile, requested_track=key)
        stem = Path(profile["tracks"][key]["outputFilename"]).stem
        paths = build_package(variant, report, OUTPUT_ROOT / "cv" / key, stem)
        results[key] = paths
        if publish:
            PUBLIC_VARIANT_DIR.mkdir(parents=True, exist_ok=True)
            shutil.copy2(paths["pdf"], PUBLIC_VARIANT_DIR / profile["tracks"][key]["publicFilename"])
            if key == "software":
                PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
                shutil.copy2(paths["pdf"], PUBLIC_DIR / MASTER_FILENAME)
    return results


def validate_pdf_is_one_page(path: Path) -> None:
    reader = PdfReader(str(path))
    if len(reader.pages) != 1:
        raise ValueError(f"Expected one page, found {len(reader.pages)}: {path}")


def private_phone() -> str | None:
    return os.environ.get("CV_PHONE", "").strip() or None
